"""Сборщик ВКР Колобаева М.Е. в DOCX по ГОСТ 7.32-2017.

Источник: Kolobaev_VKR_2026.md
Список источников: references.md
Диаграммы и графики: docs/VKR_Kolobaev/diagrams/

Запуск:
    D:\\Проекты\\uav_nav\\.venv\\Scripts\\python.exe build_docx.py

Результат: Kolobaev_VKR_2026.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).parent
PROJECT_ROOT = ROOT.parent.parent
SOURCE_MD = ROOT / "Kolobaev_VKR_2026.md"
REFERENCES_MD = ROOT / "references.md"
OUTPUT_DOCX = ROOT / "Kolobaev_VKR_2026.docx"
NBSP = " "

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14
LINE_SPACING = 1.0
MARGIN_LEFT_MM = 30
MARGIN_RIGHT_MM = 15
MARGIN_TOP_MM = 20
MARGIN_BOTTOM_MM = 20
FIRST_LINE_INDENT_CM = 1.25


def _set_run_font(run, *, bold=False, italic=False, size_pt=FONT_SIZE_PT, font_name=FONT_NAME):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)


def _set_paragraph_format(par, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True,
                          space_before_pt=0, space_after_pt=0):
    par.alignment = alignment
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM) if first_line else Cm(0)


def _setup_section(section):
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)


def _setup_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


def _add_page_number_header(section):
    header = section.header
    par = header.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    run = par.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE   \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.extend([fld_begin, instr, fld_sep, fld_end])
    _set_run_font(run)


def _clear_header(section):
    header = section.header
    par = header.paragraphs[0]
    for r in list(par.runs):
        r._element.getparent().remove(r._element)


# ------------------ Формулы OMML ------------------
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _m_var(name): return f"<m:r><m:t>{name}</m:t></m:r>"
def _m_const(text): return f"<m:r><m:rPr><m:sty m:val=\"p\"/></m:rPr><m:t>{text}</m:t></m:r>"
def _m_op(text): return f"<m:r><m:t>{text}</m:t></m:r>"
def _m_sub(b, s): return f"<m:sSub><m:e>{b}</m:e><m:sub>{s}</m:sub></m:sSub>"
def _m_sup(b, s): return f"<m:sSup><m:e>{b}</m:e><m:sup>{s}</m:sup></m:sSup>"
def _m_frac(n, d): return f"<m:f><m:num>{n}</m:num><m:den>{d}</m:den></m:f>"
def _m_sqrt(b):
    return ("<m:rad><m:radPr><m:degHide m:val=\"1\"/></m:radPr>"
            f"<m:deg/><m:e>{b}</m:e></m:rad>")
def _m_paren(b):
    return ("<m:d><m:dPr><m:begChr m:val=\"(\"/><m:endChr m:val=\")\"/></m:dPr>"
            f"<m:e>{b}</m:e></m:d>")
def _m_abs(b):
    return ("<m:d><m:dPr><m:begChr m:val=\"|\"/><m:endChr m:val=\"|\"/></m:dPr>"
            f"<m:e>{b}</m:e></m:d>")
def _m_sum(s, sup, b):
    return ("<m:nary><m:naryPr><m:chr m:val=\"∑\"/>"
            "<m:limLoc m:val=\"undOvr\"/></m:naryPr>"
            f"<m:sub>{s}</m:sub><m:sup>{sup}</m:sup><m:e>{b}</m:e></m:nary>")
def _m_omath(inner): return f"<m:oMath xmlns:m=\"{M_NS}\">{inner}</m:oMath>"


def _build_formulas():
    f = {}

    # ATE = sqrt((1/N) Σ ||p_est_i − p_gt_i||²)
    p_est = _m_sub(_m_var("p"), _m_const("est,i"))
    p_gt = _m_sub(_m_var("p"), _m_const("gt,i"))
    norm_sq = _m_sup(_m_abs(p_est + _m_op("−") + p_gt), _m_var("2"))
    f["ate"] = (
        _m_const("ATE") + _m_op("=")
        + _m_sqrt(_m_frac(_m_var("1"), _m_var("N")) + _m_sum(_m_const("i=1"), _m_var("N"), norm_sq))
    )

    # RPE = sqrt((1/M) Σ ||ΔT_est ⊖ ΔT_gt||²)
    dT_e = _m_sub(_m_var("ΔT"), _m_const("est,i"))
    dT_g = _m_sub(_m_var("ΔT"), _m_const("gt,i"))
    rpe_term = _m_sup(_m_abs(dT_e + _m_op("⊖") + dT_g), _m_var("2"))
    f["rpe"] = (
        _m_const("RPE") + _m_op("=")
        + _m_sqrt(_m_frac(_m_var("1"), _m_var("M")) + _m_sum(_m_const("i=1"), _m_var("M"), rpe_term))
    )

    # F1 = 2·P·R/(P+R)
    f["f1"] = (
        _m_const("F1") + _m_op("=")
        + _m_frac(_m_var("2") + _m_op("·") + _m_const("P") + _m_op("·") + _m_const("R"),
                  _m_const("P") + _m_op("+") + _m_const("R"))
    )

    # cosine similarity: s = (e_a · e_b) / (||e_a|| ||e_b||)
    e_a = _m_sub(_m_var("e"), _m_var("a"))
    e_b = _m_sub(_m_var("e"), _m_var("b"))
    f["cosine"] = (
        _m_var("s") + _m_op("=")
        + _m_frac(e_a + _m_op("·") + e_b, _m_abs(e_a) + _m_op("·") + _m_abs(e_b))
    )

    # Triplet Loss: L = max(0, d(a,p) − d(a,n) + α)
    d_ap = _m_const("d") + _m_paren(_m_var("a") + _m_op(",") + _m_var("p"))
    d_an = _m_const("d") + _m_paren(_m_var("a") + _m_op(",") + _m_var("n"))
    f["triplet"] = (
        _m_const("L") + _m_op("=")
        + _m_const("max") + _m_paren(_m_var("0") + _m_op(",")
                                     + d_ap + _m_op("−") + d_an + _m_op("+") + _m_var("α"))
    )

    # IoU: |A∩B| / |A∪B|
    f["iou"] = (
        _m_const("IoU") + _m_op("=")
        + _m_frac(_m_abs(_m_const("A") + _m_op("∩") + _m_const("B")),
                  _m_abs(_m_const("A") + _m_op("∪") + _m_const("B")))
    )

    # IMU preintegration: ΔR_ij = Π Exp((ω_k − b_g) Δt)
    f["preint_rot"] = (
        _m_sub(_m_var("ΔR"), _m_const("ij")) + _m_op("=")
        + _m_const("Π") + _m_const("Exp")
        + _m_paren(_m_paren(_m_sub(_m_var("ω"), _m_var("k")) + _m_op("−")
                            + _m_sub(_m_var("b"), _m_var("g"))) + _m_op("·") + _m_var("Δt"))
    )

    # EKF predict: x̂_k|k−1 = f(x̂_{k−1}, u_k)
    f["ekf_predict"] = (
        _m_sub(_m_var("x̂"), _m_const("k|k−1")) + _m_op("=")
        + _m_const("f") + _m_paren(_m_sub(_m_var("x̂"), _m_const("k−1"))
                                    + _m_op(",") + _m_sub(_m_var("u"), _m_var("k")))
    )

    # EKF update: x̂_k|k = x̂_k|k−1 + K_k (z_k − h(x̂_k|k−1))
    f["ekf_update"] = (
        _m_sub(_m_var("x̂"), _m_const("k|k")) + _m_op("=")
        + _m_sub(_m_var("x̂"), _m_const("k|k−1")) + _m_op("+") + _m_sub(_m_var("K"), _m_var("k"))
        + _m_paren(_m_sub(_m_var("z"), _m_var("k")) + _m_op("−")
                   + _m_const("h") + _m_paren(_m_sub(_m_var("x̂"), _m_const("k|k−1"))))
    )

    # Adaptive corridor: D_max = D_min + (D_cap − D_min) · min(1, ρ/ρ_ref)
    rho_ref = _m_sub(_m_var("ρ"), _m_const("ref"))
    f["corridor"] = (
        _m_sub(_m_var("D"), _m_const("max")) + _m_op("=")
        + _m_sub(_m_var("D"), _m_const("min")) + _m_op("+")
        + _m_paren(_m_sub(_m_var("D"), _m_const("cap")) + _m_op("−")
                   + _m_sub(_m_var("D"), _m_const("min")))
        + _m_op("·") + _m_const("min")
        + _m_paren(_m_var("1") + _m_op(",") + _m_frac(_m_var("ρ"), rho_ref))
    )

    # Mahalanobis: γ² = (z − h)ᵀ S⁻¹ (z − h)
    f["mahalanobis"] = (
        _m_sup(_m_var("γ"), _m_var("2")) + _m_op("=")
        + _m_sup(_m_paren(_m_var("z") + _m_op("−") + _m_const("h")), _m_const("T"))
        + _m_sup(_m_var("S"), _m_var("−1"))
        + _m_paren(_m_var("z") + _m_op("−") + _m_const("h"))
    )

    # Геометрический расчёт ΔX = V · Δt
    f["delta_x"] = (
        _m_const("ΔX") + _m_op("=") + _m_const("V") + _m_op("·") + _m_var("Δt")
    )

    # FPS = 1000 / t_inf
    f["fps"] = (
        _m_const("FPS") + _m_op("=") + _m_frac(_m_var("1000"), _m_sub(_m_var("t"), _m_const("inf")))
    )

    # mAP@τ = (1/C) Σ AP_c (τ)
    AP_c = _m_sub(_m_const("AP"), _m_var("c")) + _m_paren(_m_var("τ"))
    f["map"] = (
        _m_const("mAP") + _m_op("@") + _m_var("τ") + _m_op("=")
        + _m_frac(_m_var("1"), _m_const("C")) + _m_sum(_m_const("c=1"), _m_const("C"), AP_c)
    )

    # Inline единичные
    f["alpha_t"] = _m_sub(_m_var("ᾱ"), _m_var("t"))
    f["sigma"] = _m_var("σ")
    f["mu"] = _m_var("μ")

    return f


FORMULAS_OMML = _build_formulas()

RE_H1 = re.compile(r"^# (.+)$")
RE_H2 = re.compile(r"^## (.+)$")
RE_H3 = re.compile(r"^### (.+)$")
RE_H4 = re.compile(r"^#### (.+)$")
RE_FIG = re.compile(r"^>\s*FIG:\s*([^|]+?)\s*\|\s*(.+)$")
RE_FIG_SMALL = re.compile(r"^>\s*FIG_SMALL:\s*([^|]+?)\s*\|\s*(.+)$")
RE_TBL = re.compile(r"^>\s*TBL:\s*(.+)$")
RE_FORM = re.compile(r"^>\s*FORM:\s*([a-zA-Z0-9_]+)\s*$")
RE_TOC = re.compile(r"^>\s*TOC\s*$")
RE_TASK_PAGE = re.compile(r"^>\s*TASK_PAGE\s*$")
RE_CODE_APPENDIX = re.compile(r"^>\s*CODE_APPENDIX:\s*([^|]+?)(?:\s*\|\s*(.+))?$")
RE_REFS = re.compile(r"^>\s*REFS\s*$")
RE_INLINE_OMML = re.compile(r"\$([a-zA-Z_][a-zA-Z0-9_]*)\$")


def _parse_inline_to_runs(text):
    tokens = []
    pos = 0
    pattern = re.compile(r"(\*\*(?P<b>.+?)\*\*|\*(?P<i>[^*\n]+?)\*|`(?P<c>[^`]+)`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], {}))
        if m.group("b") is not None:
            tokens.append((m.group("b"), {"bold": True}))
        elif m.group("i") is not None:
            tokens.append((m.group("i"), {"italic": True}))
        elif m.group("c") is not None:
            tokens.append((m.group("c"), {"italic": True}))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], {}))
    return tokens


def _append_inline_omml(par, fid):
    if fid not in FORMULAS_OMML:
        run = par.add_run(f"${fid}$")
        _set_run_font(run)
        return
    omath_xml = _m_omath(FORMULAS_OMML[fid])
    omath = parse_xml(omath_xml)
    par._p.append(omath)


def _add_runs_with_formatting(par, text, *, base_bold, base_size_pt):
    for piece, flags in _parse_inline_to_runs(text):
        if not piece:
            continue
        run = par.add_run(piece)
        _set_run_font(run, bold=flags.get("bold", base_bold),
                      italic=flags.get("italic", False), size_pt=base_size_pt)


def _add_paragraph_with_inline(doc, text, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line=True, space_after_pt=0, space_before_pt=0,
                               base_bold=False, base_size_pt=FONT_SIZE_PT):
    par = doc.add_paragraph()
    _set_paragraph_format(par, alignment=alignment, first_line=first_line,
                          space_before_pt=space_before_pt, space_after_pt=space_after_pt)
    pos = 0
    for m in RE_INLINE_OMML.finditer(text):
        if m.start() > pos:
            _add_runs_with_formatting(par, text[pos:m.start()],
                                      base_bold=base_bold, base_size_pt=base_size_pt)
        _append_inline_omml(par, m.group(1))
        pos = m.end()
    if pos < len(text):
        _add_runs_with_formatting(par, text[pos:],
                                  base_bold=base_bold, base_size_pt=base_size_pt)
    return par


def _apply_outline_level(par, level):
    pPr = par._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(level))


def _add_heading_h1(doc, text):
    par = doc.add_paragraph(style="Heading 1")
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(36)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    run = par.add_run(text)
    _set_run_font(run, bold=False, size_pt=14)
    pPr = par._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))
    _apply_outline_level(par, 0)


def _add_heading_h2(doc, text):
    par = doc.add_paragraph(style="Heading 2")
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    run = par.add_run(text)
    _set_run_font(run, bold=False, size_pt=14)
    _apply_outline_level(par, 1)


def _add_heading_h3(doc, text):
    par = doc.add_paragraph(style="Heading 3")
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    run = par.add_run(text)
    _set_run_font(run, bold=False, size_pt=14)
    _apply_outline_level(par, 2)


def _add_heading_h4(doc, text):
    par = doc.add_paragraph(style="Heading 4")
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    run = par.add_run(text)
    _set_run_font(run, bold=False, size_pt=14)
    _apply_outline_level(par, 3)


def add_table_of_contents(doc):
    par = doc.add_paragraph()
    _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
                          space_before_pt=18, space_after_pt=18)
    pPr = par._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))
    run = par.add_run("Содержание")
    _set_run_font(run, bold=False, size_pt=14)

    par_toc = doc.add_paragraph()
    _set_paragraph_format(par_toc, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False,
                          space_after_pt=0)
    run = par_toc.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление наполняется при открытии в Word: ПКМ → «Обновить поле»."
    placeholder_run = OxmlElement("w:r")
    placeholder_run.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.extend([fld_begin, instr, fld_sep])
    par_toc._p.append(placeholder_run)
    run2 = par_toc.add_run()
    run2._element.append(fld_end)
    _set_run_font(run, size_pt=FONT_SIZE_PT)
    _set_run_font(run2, size_pt=FONT_SIZE_PT)


def _compute_picture_width_cm(image_path, default_cm=15.0, max_height_cm=21.5, scale=1.0):
    base = default_cm * scale
    max_h = max_height_cm * scale if scale < 1.0 else max_height_cm
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(image_path) as im:
            w, h = im.size
        if w <= 0 or h <= 0:
            return base
        ratio = w / h
        height_at_default = base / ratio
        if height_at_default <= max_h:
            return base
        return round(max_h * ratio, 2)
    except Exception:
        return base


def _add_figure(doc, image_path, caption, *, fig_no, scale=1.0):
    if not image_path.exists():
        _add_paragraph_with_inline(doc, f"[Изображение не найдено: {image_path}]",
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        return
    par_img = doc.add_paragraph()
    _set_paragraph_format(par_img, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
                          space_before_pt=12, space_after_pt=0)
    par_img.paragraph_format.keep_with_next = True
    par_img.paragraph_format.keep_together = True
    run = par_img.add_run()
    width_cm = _compute_picture_width_cm(image_path, scale=scale)
    try:
        run.add_picture(str(image_path), width=Cm(width_cm))
    except Exception:
        run.add_text(f"[Не удалось вставить {image_path.name}]")
    par_cap = doc.add_paragraph()
    _set_paragraph_format(par_cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
                          space_before_pt=0, space_after_pt=12)
    par_cap.paragraph_format.keep_together = True
    run = par_cap.add_run(f"Рисунок {fig_no} − {caption}")
    _set_run_font(run, size_pt=FONT_SIZE_PT)


# --- Титул, задание, источники, приложения, таблицы, формулы ---

def build_task_page(doc):
    def _line(text, *, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=False, space_after=0, space_before=0, first_line=False):
        p = doc.add_paragraph()
        _set_paragraph_format(p, alignment=align, first_line=first_line,
                              space_before_pt=space_before, space_after_pt=space_after)
        run = p.add_run(text)
        _set_run_font(run, bold=bold, size_pt=size)
        return p

    p = _line("Министерство науки и высшего образования Российской Федерации", space_before=0)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))
    _line("ФГБОУ ВО «Ярославский государственный технический университет»")
    _line("Кафедра «Кибернетика»", space_after=18)

    _line("УТВЕРЖДАЮ", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=12)
    _line("Заведующий кафедрой,", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _line("канд. техн. наук", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _line("________ И. В. Тюкин", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _line("«____» _______________ 2026 г.", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    _line("ЗАДАНИЕ", bold=True, size=14, space_before=6)
    _line("на выпускную квалификационную работу магистра")
    _line("студенту группы ЦМИИ-20М Колобаеву Максиму Евгеньевичу", space_after=12)

    def field(label, value):
        par = doc.add_paragraph()
        _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False,
                              space_after_pt=4)
        run = par.add_run(f"{label}: ")
        _set_run_font(run, bold=True, size_pt=12)
        run = par.add_run(value)
        _set_run_font(run, size_pt=12)

    field("1. Тема работы",
          "«Автономная визуально-инерциальная навигация беспилотного летательного аппарата "
          "с семантическим landmark-matching в условиях отсутствия GNSS-сигнала на основе датасета MidAir»")
    field("2. Срок сдачи студентом законченной работы", "« __ » июня 2026 г.")
    field("3. Исходные данные к работе",
          "синтетический датасет MidAir (университет Льежа); открытые модели YOLOv8/YOLO11, SegFormer; "
          "результаты НИР (отчёт ЯГТУ 27.04.04 — 002 НИР, 2026 г.); GPU-сервер с NVIDIA RTX 3090; "
          "одноплатный компьютер Raspberry Pi 5 (8 ГБ).")
    field("4. Содержание расчётно-пояснительной записки",
          "введение; глава 1 — анализ предметной области, обзор современных исследований и сравнительный "
          "анализ методов; глава 2 — методология и архитектура предлагаемой системы автономной навигации; "
          "глава 3 — программная реализация и испытания компонентов; глава 4 — экспериментальное исследование "
          "и анализ полученных результатов; заключение; список использованных источников; приложения.")
    field("5. Перечень графического материала",
          "C4-диаграммы архитектуры системы; UML-диаграммы классов модулей восприятия, памяти, оценивания "
          "и планирования; sequence- и activity-диаграммы алгоритмов; use case диаграмма режимов миссии; "
          "графики обучения моделей сегментации и embedding-головы; кривые ATE/RPE; графики FPS и латентности "
          "на Raspberry Pi 5; визуализации траекторий полёта.")
    field("6. Дата выдачи задания", "« __ » _____________ 2025 г.")

    _line("Руководитель      _______________________ / И. В. Тюкин /",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=18)
    _line("Задание принял к исполнению      _______________________ / М. Е. Колобаев /",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=12)


def build_title_page(doc):
    def _center(text, *, bold=False, size=FONT_SIZE_PT, space_after=0, space_before=0):
        par = doc.add_paragraph()
        _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
                              space_before_pt=space_before, space_after_pt=space_after)
        run = par.add_run(text)
        _set_run_font(run, bold=bold, size_pt=size)
        return par

    def _two_col(left_lines, right_lines, *, font_size=FONT_SIZE_PT):
        n_rows = max(len(left_lines), len(right_lines))
        if n_rows == 0:
            return
        table = doc.add_table(rows=n_rows, cols=2)
        table.autofit = False
        _disable_table_borders(table)
        col_widths = [Cm(8.0), Cm(8.5)]
        for i, w in enumerate(col_widths):
            table.columns[i].width = w
            for cell in table.columns[i].cells:
                cell.width = w
                _set_cell_no_padding(cell, top_dxa=0, bottom_dxa=0, left_dxa=0, right_dxa=0)
        for r in range(n_rows):
            for col_idx, lines in enumerate((left_lines, right_lines)):
                cell_par = table.rows[r].cells[col_idx].paragraphs[0]
                pf = cell_par.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.first_line_indent = Cm(0)
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                pf.line_spacing = 1.0
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                if r < len(lines):
                    run = cell_par.add_run(lines[r])
                    _set_run_font(run, size_pt=font_size)

    _center("Министерство науки и высшего образования Российской Федерации")
    _center("Федеральное государственное бюджетное образовательное учреждение")
    _center("высшего образования")
    _center("«Ярославский государственный технический университет»")
    _center("Кафедра «Кибернетика»", space_after=18)

    left = ["УДК 004.93+629.7"]
    right = [
        "ДОПУСКАЕТСЯ К ЗАЩИТЕ",
        "Заведующий кафедрой,",
        "канд. техн. наук",
        "________ И. В. Тюкин",
        "«__» _________ 2026 г.",
    ]
    _two_col(left, right)
    doc.add_paragraph()

    _center("АВТОНОМНАЯ ВИЗУАЛЬНО-ИНЕРЦИАЛЬНАЯ НАВИГАЦИЯ", bold=True, size=14)
    _center("БЕСПИЛОТНОГО ЛЕТАТЕЛЬНОГО АППАРАТА", bold=True, size=14)
    _center("С СЕМАНТИЧЕСКИМ LANDMARK-MATCHING В УСЛОВИЯХ", bold=True, size=14)
    _center("ОТСУТСТВИЯ GNSS-СИГНАЛА НА ОСНОВЕ ДАТАСЕТА MIDAIR",
            bold=True, size=14, space_after=12)

    _center("Пояснительная записка к выпускной квалификационной работе магистра")
    _center("по направлению 27.04.04 «Управление в технических системах»", space_after=12)
    _center("ЯГТУ 27.04.04 – 002 ВКР", bold=True)
    doc.add_paragraph()

    left = [
        "СОГЛАСОВАНО",
        "Руководитель,",
        "зав. кафедрой, канд. техн. наук",
        "________ И. В. Тюкин",
        "«__» _________ 2026 г.",
    ]
    right = [
        "",
        "Нормоконтролёр,",
        "доцент, канд. техн. наук",
        "________ И. В. Тюкин",
        "«__» _________ 2026 г.",
    ]
    _two_col(left, right)
    doc.add_paragraph()

    left = [""]
    right = [
        "Работу выполнил",
        "студент группы ЦМИИ-20М",
        "________ М. Е. Колобаев",
        "«__» _________ 2026 г.",
    ]
    _two_col(left, right)

    par = doc.add_paragraph()
    _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
                          space_before_pt=24, space_after_pt=0)
    run = par.add_run("Ярославль, 2026")
    _set_run_font(run, size_pt=FONT_SIZE_PT)


def add_references_section(doc):
    if not REFERENCES_MD.exists():
        return
    text = REFERENCES_MD.read_text(encoding="utf-8")
    items = []
    for line in text.splitlines():
        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            items.append(m.group(2))
    if not items:
        return
    _add_heading_h1(doc, "Список использованных источников")
    for idx, item in enumerate(items, start=1):
        body = re.sub(r"\s+-\s+", f"{NBSP}-{NBSP}", item)
        par = doc.add_paragraph()
        _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True,
                              space_after_pt=6)
        run = par.add_run(f"{idx}.{NBSP}")
        _set_run_font(run)
        for piece, flags in _parse_inline_to_runs(body):
            if not piece:
                continue
            run = par.add_run(piece)
            _set_run_font(run, bold=flags.get("bold", False),
                          italic=flags.get("italic", False))


def add_code_appendix(doc, file_paths, *, max_lines=700, appendix_letter="А", subtitle="Исходный код"):
    par = doc.add_paragraph(style="Heading 1")
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.keep_with_next = True
    run = par.add_run(f"Приложение {appendix_letter}")
    _set_run_font(run, bold=False, size_pt=14)
    pPr = par._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))
    _apply_outline_level(par, 0)

    par2 = doc.add_paragraph()
    pf2 = par2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf2.line_spacing = LINE_SPACING
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(24)
    pf2.first_line_indent = Cm(0)
    pf2.keep_with_next = True
    run2 = par2.add_run(subtitle)
    _set_run_font(run2, bold=False, size_pt=14)

    lines_used = 0
    for idx, file_path in enumerate(file_paths, start=1):
        if not file_path.exists():
            _add_paragraph_with_inline(doc, f"[Файл не найден: {file_path}]",
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            continue
        if lines_used >= max_lines:
            break
        try:
            rel = file_path.relative_to(PROJECT_ROOT)
        except ValueError:
            rel = file_path
        _add_heading_h2(doc, f"Листинг {appendix_letter}.{idx} - {rel.as_posix()}")
        source_lines = file_path.read_text(encoding="utf-8", errors="replace").splitlines()
        total = len(source_lines)
        remaining = max_lines - lines_used
        truncated = total > remaining
        if truncated:
            source_lines = source_lines[:remaining]
        lines_used += len(source_lines)
        body_par = doc.add_paragraph()
        pf = body_par.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)
        for li, src_line in enumerate(source_lines):
            run = body_par.add_run(src_line if src_line else " ")
            _set_run_font(run, size_pt=9, font_name="Consolas")
            if li < len(source_lines) - 1:
                br = OxmlElement("w:br")
                run._element.append(br)
        if truncated:
            note = doc.add_paragraph()
            _set_paragraph_format(note, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False,
                                  space_before_pt=2, space_after_pt=4)
            run = note.add_run(f"[... листинг сокращён, полный объём {total} строк ...]")
            _set_run_font(run, italic=True, size_pt=9)
            break


def _disable_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_no_padding(cell, *, top_dxa=0, bottom_dxa=0, left_dxa=0, right_dxa=0):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for side, value in (("top", top_dxa), ("left", left_dxa),
                        ("bottom", bottom_dxa), ("right", right_dxa)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(value))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _add_displayed_formula(doc, omml_inner, *, formula_no):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _disable_table_borders(table)
    col_widths = [Cm(1.5), Cm(13.5), Cm(1.5)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        for cell in table.columns[i].cells:
            cell.width = width
            _set_cell_no_padding(cell, top_dxa=120, bottom_dxa=120, left_dxa=0, right_dxa=0)
    row = table.rows[0]

    def _set_par(par, *, alignment):
        pf = par.paragraph_format
        pf.alignment = alignment
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    _set_par(row.cells[0].paragraphs[0], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    fml_par = row.cells[1].paragraphs[0]
    _set_par(fml_par, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    omath = parse_xml(_m_omath(omml_inner))
    fml_par._p.append(omath)
    num_par = row.cells[2].paragraphs[0]
    _set_par(num_par, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    num_run = num_par.add_run(f"({formula_no})")
    _set_run_font(num_run)


def _add_table_caption(doc, caption, *, tbl_no):
    par = doc.add_paragraph()
    _set_paragraph_format(par, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=True,
                          space_before_pt=12, space_after_pt=0)
    par.paragraph_format.keep_with_next = True
    par.paragraph_format.keep_together = True
    run = par.add_run(f"Таблица {tbl_no} − {caption}")
    _set_run_font(run, size_pt=FONT_SIZE_PT)


def _add_markdown_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell_par = cell.paragraphs[0]
            cell_par.paragraph_format.first_line_indent = Cm(0)
            cell_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pos = 0
            for m in RE_INLINE_OMML.finditer(text):
                if m.start() > pos:
                    for piece, flags in _parse_inline_to_runs(text[pos:m.start()]):
                        if not piece:
                            continue
                        run = cell_par.add_run(piece)
                        _set_run_font(run, bold=False,
                                      italic=flags.get("italic", False),
                                      size_pt=FONT_SIZE_PT - 1)
                _append_inline_omml(cell_par, m.group(1))
                pos = m.end()
            if pos < len(text):
                for piece, flags in _parse_inline_to_runs(text[pos:]):
                    if not piece:
                        continue
                    run = cell_par.add_run(piece)
                    _set_run_font(run, bold=flags.get("bold", False),
                                  italic=flags.get("italic", False),
                                  size_pt=FONT_SIZE_PT - 1)


def _parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _is_table_separator(line):
    s = line.strip().strip("|")
    return bool(s) and all(set(c.strip()).issubset(set("-:")) for c in s.split("|") if c.strip())


def render_md_to_doc(doc):
    text = SOURCE_MD.read_text(encoding="utf-8")
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = text.replace("—", "-").replace("–", "-")

    fig_no = 0
    tbl_no = 0
    formula_no = 0
    pending_caption = None
    table_rows = []
    in_table = False

    def flush_table():
        nonlocal in_table, table_rows, pending_caption
        if table_rows:
            _add_markdown_table(doc, table_rows)
            trailer = doc.add_paragraph()
            tpf = trailer.paragraph_format
            tpf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            tpf.line_spacing = Pt(12)
            tpf.space_before = Pt(0)
            tpf.space_after = Pt(0)
            tpf.first_line_indent = Cm(0)
        table_rows = []
        in_table = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if _is_table_separator(line):
                continue
            row = _parse_table_row(line)
            if not in_table:
                in_table = True
            table_rows.append(row)
            continue
        else:
            if in_table:
                flush_table()

        if not line.strip():
            continue

        if m := RE_H1.match(line):
            _add_heading_h1(doc, m.group(1)); continue
        if m := RE_H2.match(line):
            _add_heading_h2(doc, m.group(1)); continue
        if m := RE_H4.match(line):
            _add_heading_h4(doc, m.group(1)); continue
        if m := RE_H3.match(line):
            _add_heading_h3(doc, m.group(1)); continue
        if m := RE_FIG.match(line):
            fig_no += 1
            img_path = (ROOT / m.group(1).strip()).resolve()
            _add_figure(doc, img_path, m.group(2).strip(), fig_no=fig_no); continue
        if m := RE_FIG_SMALL.match(line):
            fig_no += 1
            img_path = (ROOT / m.group(1).strip()).resolve()
            _add_figure(doc, img_path, m.group(2).strip(), fig_no=fig_no, scale=0.65); continue
        if RE_TOC.match(line):
            add_table_of_contents(doc); continue
        if RE_TASK_PAGE.match(line):
            build_task_page(doc); continue
        if m := RE_CODE_APPENDIX.match(line):
            rel_paths = [p.strip() for p in m.group(1).split(",")]
            paths = [(PROJECT_ROOT / p).resolve() for p in rel_paths]
            extra = (m.group(2) or "").strip()
            args = {}
            for kv in extra.split(";"):
                kv = kv.strip()
                if "=" in kv:
                    k, v = kv.split("=", 1)
                    args[k.strip()] = v.strip()
            add_code_appendix(
                doc, paths,
                max_lines=int(args.get("max_lines", 700)),
                appendix_letter=args.get("letter", "А"),
                subtitle=args.get("title", "Исходный код"),
            ); continue
        if RE_REFS.match(line):
            add_references_section(doc); continue
        if m := RE_TBL.match(line):
            tbl_no += 1
            pending_caption = m.group(1).strip()
            _add_table_caption(doc, pending_caption, tbl_no=tbl_no); continue
        if m := RE_FORM.match(line):
            fid = m.group(1).strip()
            if fid not in FORMULAS_OMML:
                _add_paragraph_with_inline(doc, f"[Неизвестная формула: {fid}]",
                                           alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                continue
            formula_no += 1
            _add_displayed_formula(doc, FORMULAS_OMML[fid], formula_no=formula_no); continue
        if line.startswith("- "):
            _add_paragraph_with_inline(doc, "- " + line[2:].strip(),
                                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                       first_line=True, space_after_pt=0)
            continue
        if re.match(r"^\d+\.\s+", line):
            _add_paragraph_with_inline(doc, line,
                                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                       first_line=True, space_after_pt=0)
            continue
        _add_paragraph_with_inline(doc, line)

    if in_table:
        flush_table()


def main():
    doc = Document()
    _setup_default_style(doc)
    section1 = doc.sections[0]
    _setup_section(section1)
    _clear_header(section1)

    build_title_page(doc)

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    _setup_section(section2)
    section2.header.is_linked_to_previous = False
    _add_page_number_header(section2)

    render_md_to_doc(doc)

    doc.save(OUTPUT_DOCX)
    print(f"OK -> {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
